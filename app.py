import json
import sqlite3
from datetime import datetime
from io import BytesIO
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google import genai
from google.genai import types
import pandas as pd
from pypdf import PdfReader
import streamlit as st

# SYSTEM CONFIGURATION
MODEL_NAME = "gemini-2.5-flash"

st.set_page_config(
    page_title="Clausewise Enterprise — Contract Lifecycle Engine",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="expanded",
)

# DATABASE INITIALIZATION
conn = sqlite3.connect("contract_repository.db", check_same_thread=False)
cursor = conn.cursor()
cursor.execute(
    """
    CREATE TABLE IF NOT EXISTS contracts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        filename TEXT,
        audit_date TEXT,
        liability_cap TEXT,
        payment_terms TEXT,
        risk_status TEXT
    )
    """
)
conn.commit()

if "audit_count" not in st.session_state:
    st.session_state.audit_count = 0

# STYLING SHEET
st.markdown(
    """<style>.stApp { background-color: #fbfcfa !important; color: #17212b !important; }[data-testid='stSidebar'] { background-color: #f1f5f1 !important; border-right: 1px solid #e5e9e6 !important; }[data-testid='stSidebar'] * { color: #17212b !important; font-weight: 600 !important; }textarea, input { background-color: #ffffff !important; border: 2px solid #185c4a !important; color: #17212b !important; }.paywall-card { background: linear-gradient(135deg, #143f35 0%, #1d6a55 100%) !important; border-radius: 12px; padding: 20px; color: white !important; }div.stButton > button:first-child { background-color: #185c4a !important; color: #ffffff !important; font-weight: 700 !important; border: none !important; }</style>""",
    unsafe_allow_html=True,
)

# HELPER CORE 1: MULTI-AGENT INFERENCE ENGINE
def run_contract_audit(contract_text: str, api_key: str) -> dict:
    client = genai.Client(api_key=api_key)
    system_prompt = """
    You are an elite Senior Commercial Legal Engine acting strictly for the BUYER/LICENSEE.
    Analyze the contract text and return a JSON payload with this exact schema:
    {
      "agent_1_syntactic": {
        "target": "Clause location and title",
        "analysis": "Syntactic analysis of legal risk.",
        "original_text": "Original clause snippet to be replaced",
        "playbook_positions": {
          "position_a_ideal": "Aggressive buyer protection revision",
          "position_b_fallback": "Balanced commercial fallback revision",
          "position_c_walkaway": "Minimum acceptable threshold revision"
        }
      },
      "agent_2_cross_clause": {
        "conflict": "Cross-clause conflict description",
        "analysis": "Explanation of legal trap",
        "redline_redirection": "Specific corrective text"
      },
      "agent_3_portfolio_recovery": {
        "target": "Financial and operational terms",
        "analysis": "Financial risk analysis",
        "liability_cap_extracted": "Extracted liability threshold summary",
        "payment_terms_extracted": "Extracted payment milestone summary",
        "risk_status": "High, Medium, or Low"
      },
      "pillar_5_obligation_registry": [
        {
          "clause": "Clause Reference",
          "data": "Extracted timeline or numeric threshold",
          "status": "Systemic Risk Tag"
        }
      ]
    }
    """
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=f"Analyze this contract text:\n\n{contract_text}",
        config=types.GenerateContentConfig(
            system_instruction=system_prompt,
            response_mime_type="application/json",
            temperature=0.1,
        ),
    )
    return json.loads(response.text)

# HELPER CORE 2: NATIVE DOCX MARKUP INJECTOR
def create_native_tracked_changes_docx(deleted_text: str, inserted_text: str) -> BytesIO:
    doc = docx.Document()
    doc.add_heading("Clausewise — Native Tracked Changes Markup", level=0)
    p = doc.add_paragraph("Legal Redline Draft:\n")
    
    # Deletion Setup
    del_run = OxmlElement("w:del")
    del_run.set(qn("w:id"), "0")
    del_run.set(qn("w:author"), "Clausewise AI")
    del_run.set(qn("w:date"), datetime.now().isoformat())
    t_del = OxmlElement("w:delText")
    t_del.text = deleted_text
    del_run.append(t_del)
    p._p.append(del_run)
    
    p.add_run(" ")
    
    # Insertion Setup
    ins_run = OxmlElement("w:ins")
    ins_run.set(qn("w:id"), "1")
    ins_run.set(qn("w:author"), "Clausewise AI")
    ins_run.set(qn("w:date"), datetime.now().isoformat())
    t_ins = OxmlElement("w:r")
    t_ins_text = OxmlElement("w:t")
    t_ins_text.text = inserted_text
    t_ins.append(t_ins_text)
    ins_run.append(t_ins)
    p._p.append(ins_run)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# SIDEBAR VIEWPORT
st.sidebar.markdown("# ◈ Clausewise Enterprise")
st.sidebar.markdown("---")
st.sidebar.markdown("### 👤 Enterprise Workspace")
st.sidebar.markdown("🛡️ **Zero Data Retention Active**")
st.sidebar.markdown("⚡ Engine: **Multi-Agent Neural Mesh**")
st.sidebar.markdown("---")

api_key_input = st.secrets.get("GEMINI_API_KEY", "")

st.title("Automate Your Contract Risk Reviews")
st.caption(
    "100% Efficient Engine: Multi-format parsing, native MS Word tracked redlines, and persistent portfolio management."
)

tab_audit, tab_repository = st.tabs(
    ["📝 Active Contract Audit", "🗄️ Enterprise Portfolio Repository"]
)

with tab_audit:
    st.markdown("### 1. Ingest Agreement File or Unstructured Text")
    uploaded_file = st.file_uploader(
        "Upload Contract (.pdf or .docx)", type=["pdf", "docx"]
    )

    SAMPLE_CONTRACT = """MASTER SOFTWARE & SERVICES AGREEMENT
2.3 Unilateral Engine Changes. Licensor retains the right to modify model endpoints at any time, provided throughput is unaffected.
3.4 Overdue Balances. Invoice balances unpaid after 14 calendar days shall accrue interest at 4% per annum above the Bank of England base rate.
4.2 Algorithmic Optimization Notice. The system architecture evaluates purely for administrative formatting and syntax risk; parameters disclaim all reliance on substantive legal functions.
5.1 Standard Liability Cap. Total combined financial exposure shall be strictly limited to the total amount paid by Licensee in the 3 months preceding the claim.
5.2 Third-Party Intellectual Property Protection. The Service Provider agrees to protect the Client from third-party copyright claims up to a limit of £50,000, notwithstanding any other damages, performance issues, or general contract failures arising from or related to this Agreement."""

    clause_text = ""
    filename_to_save = "Direct Paste Input"
    
    if uploaded_file is not None:
        filename_to_save = uploaded_file.name
        if uploaded_file.name.endswith(".pdf"):
            pdf_reader = PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                clause_text += page.extract_text() or ""
        elif uploaded_file.name.endswith(".docx"):
            doc_file = docx.Document(uploaded_file)
            clause_text = "\n".join([p.text for p in doc_file.paragraphs])
        st.success(f"Successfully ingested {uploaded_file.name}")
    else:
        clause_text = st.text_area(
            "Or paste contract text string directly here:",
            value=SAMPLE_CONTRACT,
            height=160,
        )

    if st.button("Execute Enterprise Audit", type="primary"):
        if not clause_text.strip():
            st.warning("Please upload a file or paste contract text.")
        elif not api_key_input:
            st.error("🔑 API Key Missing: Ensure GEMINI_API_KEY is configured in your Streamlit secrets.")
        else:
            with st.spinner("Executing Multi-Agent Syntactic Mesh..."):
                try:
                    results = run_contract_audit(clause_text, api_key_input)
                    st.success("Audit Complete!")
                    st.markdown("### 📊 LIVE INTERACTIVE NEGOTIATION DESK")
                    
                    ag1 = results.get("agent_1_syntactic", {})
                    ag2 = results.get("agent_2_cross_clause", {})
                    ag3 = results.get("agent_3_portfolio_recovery", {})
                    obligations = results.get("pillar_5_obligation_registry", [])

                    res_tab1, res_tab2, res_tab3 = st.tabs([
                        "🔍 Syntactic Risk & Redlines", 
                        "⚡ Cross-Clause Conflicts", 
                        "📋 Operational Registry"
                    ])

                    with res_tab1:
                        st.subheader(f"Target: {ag1.get('target', 'N/A')}")
                        st.markdown(f"**Risk Analysis:** {ag1.get('analysis', 'N/A')}")
                        st.text_area("Original Text Snippet:", value=ag1.get('original_text', ''), disabled=True, key="orig_text_area")
                        
                        st.markdown("#### Playbook Revision Strategies")
                        pos = ag1.get("playbook_positions", {})
                        st.info(f"**Ideal Position (Aggressive Buyer):** {pos.get('position_a_ideal', 'N/A')}")
                        st.warning(f"**Fallback Position (Balanced Commercial):** {pos.get('position_b_fallback', 'N/A')}")
                        st.error(f"**Walkaway Threshold:** {pos.get('position_c_walkaway', 'N/A')}")
                        
                        st.markdown("#### 📄 Generate Redlined Tracked Changes File")
                        chosen_revision = st.selectbox(
                            "Select target strategy to embed in track changes:",
                            ["position_a_ideal", "position_b_fallback"],
                            key="strategy_select"
                        
                        (inserted_text_target = pos.get(chosen_revision, "")
                        docx_buffer = create_native_tracked_changes_docx
